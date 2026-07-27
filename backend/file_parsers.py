"""
Universal File Parser — Extracts text and metadata from PDF, DOCX, MD, TXT, JSON, CSV files.
"""
import io
import json
import csv
from pathlib import Path
from typing import List, Dict, Any


def parse_pdf(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Extracts text from multi-page PDF documents using pypdf."""
    import pypdf
    pdf_file = io.BytesIO(file_bytes)
    reader = pypdf.PdfReader(pdf_file)
    
    documents = []
    full_text = []
    
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            full_text.append(text)
            
    doc_id = Path(filename).stem
    combined_text = "\n\n".join(full_text)
    
    if combined_text.strip():
        documents.append({
            "id": doc_id,
            "title": doc_id.replace("_", " ").replace("-", " ").title(),
            "abstract": combined_text,
            "file_type": "pdf",
            "page_count": len(reader.pages)
        })
    return documents


def parse_docx(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Extracts text from Word DOCX documents using python-docx."""
    import docx
    docx_file = io.BytesIO(file_bytes)
    doc = docx.Document(docx_file)
    
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    # Also extract text inside tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
                
    doc_id = Path(filename).stem
    combined_text = "\n\n".join(paragraphs)
    
    documents = []
    if combined_text.strip():
        documents.append({
            "id": doc_id,
            "title": doc_id.replace("_", " ").replace("-", " ").title(),
            "abstract": combined_text,
            "file_type": "docx",
            "paragraph_count": len(paragraphs)
        })
    return documents


def parse_text_markdown(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Extracts text from Markdown (.md) and Plain Text (.txt) files."""
    text = file_bytes.decode("utf-8", errors="ignore")
    doc_id = Path(filename).stem
    
    documents = []
    if text.strip():
        documents.append({
            "id": doc_id,
            "title": doc_id.replace("_", " ").replace("-", " ").title(),
            "abstract": text,
            "file_type": Path(filename).suffix.lstrip("."),
        })
    return documents


def parse_json(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Parses JSON files containing lists of items, FAQs, or documents."""
    raw_str = file_bytes.decode("utf-8", errors="ignore")
    data = json.loads(raw_str)
    
    documents = []
    base_id = Path(filename).stem
    
    if isinstance(data, list):
        for idx, item in enumerate(data):
            if isinstance(item, dict):
                text = item.get("text") or item.get("abstract") or item.get("content") or item.get("answer") or str(item)
                title = item.get("title") or item.get("question") or f"{base_id}_item_{idx+1}"
                doc_id = item.get("id") or f"{base_id}_{idx+1}"
                documents.append({
                    "id": str(doc_id),
                    "title": str(title),
                    "abstract": str(text),
                    "file_type": "json"
                })
            else:
                documents.append({
                    "id": f"{base_id}_{idx+1}",
                    "title": f"{base_id} Item {idx+1}",
                    "abstract": str(item),
                    "file_type": "json"
                })
    elif isinstance(data, dict):
        text = data.get("text") or data.get("content") or str(data)
        title = data.get("title") or base_id
        documents.append({
            "id": base_id,
            "title": str(title),
            "abstract": str(text),
            "file_type": "json"
        })
    return documents


def parse_csv(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Parses CSV rows into individual documents."""
    text_content = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.DictReader(io.StringIO(text_content))
    
    documents = []
    base_id = Path(filename).stem
    
    for idx, row in enumerate(reader, start=1):
        row_str = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
        if row_str.strip():
            doc_id = row.get("id") or f"{base_id}_row_{idx}"
            title = row.get("title") or row.get("question") or row.get("name") or f"{base_id} Row {idx}"
            documents.append({
                "id": str(doc_id),
                "title": str(title),
                "abstract": row_str,
                "file_type": "csv"
            })
    return documents


def parse_file(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """Universal router dispatching file bytes to the appropriate format parser."""
    ext = Path(filename).suffix.lower()
    
    if ext == ".pdf":
        return parse_pdf(file_bytes, filename)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_bytes, filename)
    elif ext in [".json"]:
        return parse_json(file_bytes, filename)
    elif ext in [".csv"]:
        return parse_csv(file_bytes, filename)
    else:
        # Default fallback for .md, .txt, .py, .html, .log, etc.
        return parse_text_markdown(file_bytes, filename)


if __name__ == "__main__":
    sample_txt = b"Artificial intelligence and neural networks enhance modern information retrieval."
    docs = parse_file(sample_txt, "ai_notes.txt")
    print("Parsed sample docs:", docs)
